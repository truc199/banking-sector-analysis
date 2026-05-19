"""
Convert Markdown files to formatted DOCX
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a horizontal line via bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue
        
        # Code block
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\r\n'))
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Add shading to paragraph
            pPr = p._p.get_or_add_pPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
            pPr.append(shading)
            continue
        
        # Headings
        if line.startswith('# ') and not line.startswith('##'):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue
        
        # Blockquote
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            text = line[2:].strip()
            # Handle **bold** within blockquote
            add_formatted_text(p, text)
            # Add left border styling via shading
            pPr = p._p.get_or_add_pPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4FF"/>')
            pPr.append(shading)
            i += 1
            continue
        
        # Table
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Filter out separator lines
            data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|$', l)]
            
            if len(data_lines) < 1:
                continue
            
            # Parse cells
            rows_data = []
            for tl in data_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                rows_data.append(cells)
            
            if not rows_data:
                continue
            
            num_cols = max(len(r) for r in rows_data)
            
            # Create table
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for ri, row_data in enumerate(rows_data):
                for ci in range(num_cols):
                    cell = table.cell(ri, ci)
                    cell_text = row_data[ci] if ci < len(row_data) else ''
                    
                    # Clear default paragraph
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    
                    add_formatted_text(p, cell_text, default_size=Pt(9))
                    
                    # Header row styling
                    if ri == 0:
                        set_cell_shading(cell, '1F4E79')
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                            run.font.size = Pt(9)
                    else:
                        # Alternate row shading
                        if ri % 2 == 0:
                            set_cell_shading(cell, 'F2F7FB')
            
            doc.add_paragraph()  # spacing after table
            continue
        
        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line.strip())
        i += 1
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def add_formatted_text(paragraph, text, default_size=None):
    """Parse markdown bold/italic and add runs to paragraph."""
    # Split by **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        
        run.font.name = 'Calibri'
        if default_size:
            run.font.size = default_size


if __name__ == '__main__':
    md_to_docx(
        r'd:\uni\gcontest\storytelling_flow copy 3.md',
        r'd:\uni\gcontest\storytelling_flow.docx'
    )
    md_to_docx(
        r'd:\uni\gcontest\hypotheses_phase_analysis.md',
        r'd:\uni\gcontest\hypotheses_phase_analysis.docx'
    )
